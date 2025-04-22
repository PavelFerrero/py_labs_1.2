from guizero import App, Text, TextBox, PushButton, Combo, info
from materials.wallpaper import calculate_wallpaper
from materials.tile import calculate_tile
from materials.laminate import calculate_laminate
from docx import *

def calculate():
    try:
        area = float(area_input.value)
        material = material_select.value
        
        if material == "Обои":
            qty, cost = calculate_wallpaper(area)
        elif material == "Плитка":
            qty, cost = calculate_tile(area)
        elif material == "Ламинат":
            qty, cost = calculate_laminate(area)
            
        result_text.value = f"Необходимо: {qty} упаковок\nСтоимость: {cost} руб."
        
        # Сохранение в Word
        doc = Document()
        doc.add_heading('Расчёт материалов', 0)
        doc.add_paragraph(f"Материал: {material}")
        doc.add_paragraph(f"Площадь: {area} м²")
        doc.add_paragraph(f"Количество: {qty} упаковок")
        doc.add_paragraph(f"Общая стоимость: {cost} руб.")
        doc.save('расчёт_материалов.docx')
        
        info("Сохранено", "Результаты сохранены в файл 'расчёт_материалов.docx'")
        
    except ValueError:
        result_text.value = "Ошибка! Введите корректную площадь."

app = App(title="Калькулятор материалов", width=400, height=300)

Text(app, text="Выберите материал:")
material_select = Combo(app, options=["Обои", "Плитка", "Ламинат"])

Text(app, text="Введите площадь (м²):")
area_input = TextBox(app)

calculate_button = PushButton(app, text="Рассчитать", command=calculate)

result_text = Text(app, text="")

app.display()